import os
from pathlib import Path
from typing import Dict, Any, Tuple
import pypdf
import docx

class DocumentParser:
    """
    Service for parsing text and metadata from PDF, DOCX, and TXT files.
    """
    
    @staticmethod
    def parse_txt(file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parses simple text files."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            
            metadata = {
                "file_type": "txt",
                "char_count": len(content),
                "word_count": len(content.split()),
                "lines_count": len(content.splitlines())
            }
            return content, metadata
        except Exception as e:
            raise ValueError(f"Failed to parse TXT file: {str(e)}")

    @staticmethod
    def parse_pdf(file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parses PDF documents page by page."""
        try:
            reader = pypdf.PdfReader(str(file_path))
            pages_text = []
            pages_meta = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages_text.append(text)
                pages_meta.append({
                    "page_number": page_num + 1,
                    "char_count": len(text)
                })
            
            full_text = "\n\n--- Page Break ---\n\n".join(pages_text)
            
            metadata = {
                "file_type": "pdf",
                "total_pages": len(reader.pages),
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
                "pages": pages_meta
            }
            return full_text, metadata
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file: {str(e)}")

    @staticmethod
    def parse_docx(file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Parses DOCX documents (paragraphs and tables)."""
        try:
            doc = docx.Document(str(file_path))
            paragraphs_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs_text.append(para.text)
                    
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs_text.append(" | ".join(row_text))
            
            full_text = "\n\n".join(paragraphs_text)
            
            metadata = {
                "file_type": "docx",
                "char_count": len(full_text),
                "word_count": len(full_text.split()),
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables)
            }
            return full_text, metadata
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")

    @classmethod
    def parse(cls, file_path: str | Path) -> Tuple[str, Dict[str, Any]]:
        """
        Unified parser router based on file extension.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = path.suffix.lower()
        
        # General file info
        stat = path.stat()
        base_metadata = {
            "filename": path.name,
            "file_size": stat.st_size,
        }
        
        if ext == ".txt":
            content, parsed_meta = cls.parse_txt(path)
        elif ext == ".pdf":
            content, parsed_meta = cls.parse_pdf(path)
        elif ext == ".docx":
            content, parsed_meta = cls.parse_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
            
        # Combine metadatas
        combined_metadata = {**base_metadata, **parsed_meta}
        return content, combined_metadata
